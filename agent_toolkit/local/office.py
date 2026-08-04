"""Офисные инструменты: создание и инспекция .docx, .xlsx с профессиональным форматированием.

Использует python-docx / openpyxl при наличии, иначе fallback на zipfile + XML.
"""
from __future__ import annotations

import json
import re
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from typing import Any

from ..core import Tool, ToolError, Workspace

# Проверка наличия опциональных библиотек
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    _HAS_OPENPYXL = True
except ImportError:
    _HAS_OPENPYXL = False

try:
    import docx as python_docx
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_PYTHON_DOCX = True
except ImportError:
    _HAS_PYTHON_DOCX = False

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _esc(value: Any) -> str:
    text = "" if value is None else str(value)
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# ============================================================
# XLSX: openpyxl с профессиональным форматированием
# ============================================================
def _write_xlsx_openpyxl(
    path: Path,
    sheet_name: str,
    headers: list[str],
    rows: list[list[Any]],
    title: str = "",
    zebra_striping: bool = True,
    auto_width: bool = True,
) -> dict[str, Any]:
    """Создание .xlsx с форматированием через openpyxl."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = (sheet_name or "Лист1")[:31]

    # Стили
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="B4C6E7"),
        right=Side(style="thin", color="B4C6E7"),
        top=Side(style="thin", color="B4C6E7"),
        bottom=Side(style="thin", color="B4C6E7"),
    )
    zebra_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    data_font = Font(name="Calibri", size=11)

    start_row = 1

    # Заголовок отчёта (если указан)
    if title:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(headers), 1))
        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = Font(name="Calibri", size=14, bold=True, color="2F5496")
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        start_row = 3  # пустая строка после заголовка

    # Заголовки таблицы
    if headers:
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=start_row, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        data_start = start_row + 1
    else:
        data_start = start_row

    # Данные
    for row_idx, row_data in enumerate(rows):
        excel_row = data_start + row_idx
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=excel_row, column=col_idx)
            # Автоопределение типа данных
            if isinstance(value, (int, float)):
                cell.value = value
                cell.number_format = '#,##0.00' if isinstance(value, float) else '#,##0'
            else:
                cell.value = str(value) if value is not None else ""
            cell.font = data_font
            cell.border = thin_border
            cell.alignment = Alignment(vertical="center")
            # Zebra striping
            if zebra_striping and row_idx % 2 == 1:
                cell.fill = zebra_fill

    # Автоширина колонок
    if auto_width:
        for col_idx in range(1, max(len(headers), 1) + 1):
            max_len = 0
            col_letter = get_column_letter(col_idx)
            for row in ws.iter_rows(min_col=col_idx, max_col=col_idx, min_row=1, max_row=ws.max_row):
                for cell in row:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max(max_len + 4, 10), 50)

    # Закрепление заголовков
    if headers:
        ws.freeze_panes = ws.cell(row=data_start, column=1)

    # Автофильтр
    if headers and rows:
        last_col = get_column_letter(len(headers))
        ws.auto_filter.ref = f"A{start_row}:{last_col}{data_start + len(rows) - 1}"

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))

    return {
        "rows": len(rows),
        "columns": len(headers),
        "has_title": bool(title),
        "has_formatting": True,
    }


# ============================================================
# XLSX: fallback на stdlib (без форматирования)
# ============================================================
def _write_xlsx_stdlib(
    path: Path, sheet_name: str, headers: list[str], rows: list[list[Any]]
) -> dict[str, Any]:
    """Fallback: создание .xlsx через zipfile + XML."""
    sheet_name = (sheet_name or "Sheet1")[:31]
    rows_xml: list[str] = []

    def make_row_xml(row_idx: int, cells: list[Any]) -> str:
        c_xml: list[str] = []
        for i, val in enumerate(cells):
            col_letter = chr(65 + (i % 26))
            ref = f"{col_letter}{row_idx}"
            sval = _esc(val)
            c_xml.append(f'<c r="{ref}" t="inlineStr"><is><t>{sval}</t></is></c>')
        return f'<row r="{row_idx}">{"".join(c_xml)}</row>'

    if headers:
        rows_xml.append(make_row_xml(1, headers))
        start_row = 2
    else:
        start_row = 1

    for r_idx, r_cells in enumerate(rows, start=start_row):
        rows_xml.append(make_row_xml(r_idx, r_cells))

    sheet_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">\n'
        f'<sheetData>{"".join(rows_xml)}</sheetData>\n'
        "</worksheet>"
    )
    wb_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">\n'
        "<sheets>\n"
        f'  <sheet name="{_esc(sheet_name)}" sheetId="1" r:id="rId1"/>\n'
        "</sheets>\n"
        "</workbook>"
    )
    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">\n'
        '  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>\n'
        '  <Default Extension="xml" ContentType="application/xml"/>\n'
        '  <Override PartName="/xl/workbook.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>\n'
        '  <Override PartName="/xl/worksheets/sheet1.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>\n'
        "</Types>"
    )
    root_rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">\n'
        '  <Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/>\n'
        "</Relationships>"
    )
    xl_rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">\n'
        '  <Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
        'Target="worksheets/sheet1.xml"/>\n'
        "</Relationships>"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types_xml)
        z.writestr("_rels/.rels", root_rels_xml)
        z.writestr("xl/workbook.xml", wb_xml)
        z.writestr("xl/_rels/workbook.xml.rels", xl_rels_xml)
        z.writestr("xl/worksheets/sheet1.xml", sheet_xml)
    return {"rows": len(rows), "columns": len(headers), "has_formatting": False}


# ============================================================
# DOCX: python-docx с профессиональным форматированием
# ============================================================
def _write_docx_python_docx(path: Path, title: str, content: str) -> dict[str, Any]:
    """Создание .docx с форматированием через python-docx."""
    doc = python_docx.Document()

    # Настройка стилей
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    stats = {"headings": 0, "paragraphs": 0, "lists": 0, "bold": 0, "italic": 0}

    if title:
        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        # Заголовки
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            stats["headings"] += 1
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            stats["headings"] += 1
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            stats["headings"] += 1
        # Маркированный список
        elif stripped.startswith("- ") or stripped.startswith("• "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            _apply_inline_formatting(p, stripped[2:])
            stats["lists"] += 1
        # Нумерованный список
        elif re.match(r'^\d+[\.\)]\s', stripped):
            text = re.sub(r'^\d+[\.\)]\s', '', stripped)
            p = doc.add_paragraph(text, style="List Number")
            _apply_inline_formatting(p, text)
            stats["lists"] += 1
        # Горизонтальная линия
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xB4, 0xC6, 0xE7)
        # Обычный параграф с inline-форматированием
        else:
            p = doc.add_paragraph()
            _apply_inline_formatting(p, stripped)
            stats["paragraphs"] += 1

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return stats


def _apply_inline_formatting(paragraph, text: str) -> None:
    """Применить inline-форматирование: **bold**, *italic*, `code`."""
    # Очищаем paragraph от дефолтного текста
    paragraph.clear()

    # Паттерн для **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for match in pattern.finditer(text):
        # Текст до форматирования
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        pos = match.end()

    # Остаток текста
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ============================================================
# DOCX: fallback на stdlib
# ============================================================
def _write_docx_stdlib(path: Path, title: str, paragraphs: list[str]) -> dict[str, Any]:
    """Fallback: создание .docx через zipfile + XML."""
    body_xml: list[str] = []
    if title:
        body_xml.append(
            f'<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
            f'<w:r><w:t>{_esc(title)}</w:t></w:r></w:p>'
        )
    for p_text in paragraphs:
        if p_text.startswith("# "):
            h = p_text[2:].strip()
            body_xml.append(
                f'<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
                f'<w:r><w:t>{_esc(h)}</w:t></w:r></w:p>'
            )
        elif p_text.startswith("## "):
            h = p_text[3:].strip()
            body_xml.append(
                f'<w:p><w:pPr><w:pStyle w:val="Heading2"/></w:pPr>'
                f'<w:r><w:t>{_esc(h)}</w:t></w:r></w:p>'
            )
        else:
            body_xml.append(f'<w:p><w:r><w:t>{_esc(p_text)}</w:t></w:r></w:p>')

    doc_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        f'<w:document xmlns:w="{W_NS}">\n'
        f"<w:body>{''.join(body_xml)}</w:body>\n"
        "</w:document>"
    )
    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">\n'
        '  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>\n'
        '  <Default Extension="xml" ContentType="application/xml"/>\n'
        '  <Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>\n'
        "</Types>"
    )
    rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">\n'
        '  <Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>\n'
        "</Relationships>"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types_xml)
        z.writestr("_rels/.rels", rels_xml)
        z.writestr("word/document.xml", doc_xml)
    return {"headings": 0, "paragraphs": len(paragraphs), "has_formatting": False}


def build_office_tools(ws: Workspace) -> list[Tool]:
    """Собрать инструменты создания и чтения офисных документов."""

    def create_docx(path: str, title: str = "", content: str = "") -> str:
        p = ws.resolve(path)
        if p.suffix.lower() != ".docx":
            p = p.with_suffix(".docx")

        if _HAS_PYTHON_DOCX:
            stats = _write_docx_python_docx(p, title=title, content=content)
            fmt_info = (
                f", форматирование: заголовков={stats['headings']}, "
                f"параграфов={stats['paragraphs']}, списков={stats['lists']}"
            )
        else:
            paragraphs = [ln.strip() for ln in content.splitlines() if ln.strip()]
            stats = _write_docx_stdlib(p, title=title, paragraphs=paragraphs)
            fmt_info = " (без расширенного форматирования — установите python-docx)"

        return (
            f"Создан документ {ws.relative(p)} "
            f"(заголовок: {title!r}{fmt_info})"
        )

    def create_xlsx(
        path: str,
        sheet_name: str = "Лист1",
        headers_json: str = "[]",
        rows_json: str = "[]",
        title: str = "",
    ) -> str:
        p = ws.resolve(path)
        if p.suffix.lower() != ".xlsx":
            p = p.with_suffix(".xlsx")
        try:
            headers = json.loads(headers_json) if headers_json else []
            rows = json.loads(rows_json) if rows_json else []
        except ValueError as exc:
            raise ToolError(f"Некорректный JSON для headers/rows: {exc}") from exc

        if _HAS_OPENPYXL:
            info = _write_xlsx_openpyxl(p, sheet_name=sheet_name, headers=headers, rows=rows, title=title)
            fmt_info = (
                f", форматирование: заголовки (синий фон, белый шрифт), "
                f"zebra striping, автоширина, автофильтр, закреплённые строки"
            )
        else:
            info = _write_xlsx_stdlib(p, sheet_name=sheet_name, headers=headers, rows=rows)
            fmt_info = " (без форматирования — установите openpyxl)"

        return (
            f"Создана таблица {ws.relative(p)} "
            f"(лист: {sheet_name!r}, колонок: {info['columns']}, строк: {info['rows']}{fmt_info})"
        )

    def inspect_docx(path: str) -> str:
        p = ws.resolve(path)
        if not p.exists():
            raise ToolError(f"Файл {path!r} не найден")

        if _HAS_PYTHON_DOCX:
            doc = python_docx.Document(str(p))
            lines = []
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else ""
                text = para.text.strip()
                if not text:
                    continue
                if "Heading" in style_name:
                    level = style_name.replace("Heading ", "")
                    prefix = "#" * int(level) + " " if level.isdigit() else "# "
                    lines.append(prefix + text)
                else:
                    lines.append(text)
            return "\n".join(lines) if lines else "(документ пуст)"

        # Fallback
        try:
            with zipfile.ZipFile(p, "r") as z:
                xml_data = z.read("word/document.xml")
        except (zipfile.BadZipFile, KeyError, OSError) as exc:
            raise ToolError(f"Не удалось прочитать .docx {path!r}: {exc}") from exc
        root = ET.fromstring(xml_data)
        texts: list[str] = []
        for elem in root.iter():
            if elem.tag.endswith("}t") and elem.text:
                texts.append(elem.text)
        return "\n".join(texts) if texts else "(документ пуст)"

    def inspect_xlsx(path: str) -> str:
        p = ws.resolve(path)
        if not p.exists():
            raise ToolError(f"Файл {path!r} не найден")

        if _HAS_OPENPYXL:
            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws_sheet = wb[sheet]
                if len(wb.sheetnames) > 1:
                    lines.append(f"=== Лист: {sheet} ===")
                for row in ws_sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        lines.append("\t".join(cells))
            return "\n".join(lines) if lines else "(таблица пуста)"

        # Fallback
        try:
            with zipfile.ZipFile(p, "r") as z:
                sheet_xml = z.read("xl/worksheets/sheet1.xml")
        except (zipfile.BadZipFile, KeyError, OSError) as exc:
            raise ToolError(f"Не удалось прочитать .xlsx {path!r}: {exc}") from exc
        root = ET.fromstring(sheet_xml)
        rows_data: list[list[str]] = []
        for row_elem in root.iter():
            if row_elem.tag.endswith("}row"):
                cells: list[str] = []
                for cell_elem in row_elem:
                    if cell_elem.tag.endswith("}c"):
                        val = ""
                        for child in cell_elem.iter():
                            if child.tag.endswith("}t") and child.text:
                                val = child.text
                                break
                        cells.append(val)
                if cells:
                    rows_data.append(cells)
        return "\n".join("\t".join(c for c in r) for r in rows_data) if rows_data else "(таблица пуста)"

    return [
        Tool(
            name="office.create_docx",
            description="Создать документ Word (.docx) с форматированием: заголовки H1-H3, **жирный**, *курсив*, `код`, маркированные и нумерованные списки.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Путь к файлу .docx"},
                    "title": {"type": "string", "description": "Заголовок документа"},
                    "content": {
                        "type": "string",
                        "description": "Текст в Markdown (# заголовки, - списки, **bold**, *italic*, `code`)",
                    },
                },
                "required": ["path", "content"],
            },
            fn=create_docx,
            skills=["office", "documentation", "reports", "local", "docx", "word"],
            attributes={
                "category": "office",
                "read_only": False,
                "dangerous": False,
                "resource_type": "document",
                "speed": "fast",
                "tags": ["office", "docx", "word", "document", "report"],
            },
            example='office.create_docx(path="report.docx", title="Отчёт", content="## Итоги\\n- **Пункт 1**: выполнено\\n- *Пункт 2*: в процессе")',
        ),
        Tool(
            name="office.create_xlsx",
            description="Создать таблицу Excel (.xlsx) с форматированием: цветные заголовки, zebra striping, автоширина колонок, автофильтр, закреплённые строки.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Путь к файлу .xlsx"},
                    "sheet_name": {"type": "string", "description": "Название листа"},
                    "headers_json": {
                        "type": "string",
                        "description": 'JSON-массив заголовков (\'["Название", "Цена"]\')',
                    },
                    "rows_json": {
                        "type": "string",
                        "description": 'JSON-массив строк (\'[["Товар 1", 100], ["Товар 2", 200]]\')',
                    },
                    "title": {"type": "string", "description": "Опциональный заголовок отчёта над таблицей"},
                },
                "required": ["path", "headers_json", "rows_json"],
            },
            fn=create_xlsx,
            skills=["office", "documentation", "reports", "local", "xlsx", "excel"],
            attributes={
                "category": "office",
                "read_only": False,
                "dangerous": False,
                "resource_type": "spreadsheet",
                "speed": "fast",
                "tags": ["office", "xlsx", "excel", "spreadsheet", "table"],
            },
            example='office.create_xlsx(path="data.xlsx", title="Продажи", headers_json=\'["Товар", "Цена"]\', rows_json=\'[["Хлеб", 50]]\')',
        ),
        Tool(
            name="office.inspect_docx",
            description="Прочитать и извлечь текст из файла .docx с сохранением структуры заголовков.",
            parameters={
                "type": "object",
                "properties": {"path": {"type": "string", "description": "Путь к файлу .docx"}},
                "required": ["path"],
            },
            fn=inspect_docx,
            skills=["office", "documentation", "local", "docx", "read"],
            attributes={
                "category": "office", "read_only": True, "dangerous": False,
                "resource_type": "document", "speed": "fast",
                "tags": ["office", "docx", "word", "read", "text"],
            },
            example='office.inspect_docx(path="report.docx")',
        ),
        Tool(
            name="office.inspect_xlsx",
            description="Прочитать все листы и ячейки таблицы .xlsx в виде текста.",
            parameters={
                "type": "object",
                "properties": {"path": {"type": "string", "description": "Путь к файлу .xlsx"}},
                "required": ["path"],
            },
            fn=inspect_xlsx,
            skills=["office", "documentation", "local", "xlsx", "read"],
            attributes={
                "category": "office", "read_only": True, "dangerous": False,
                "resource_type": "spreadsheet", "speed": "fast",
                "tags": ["office", "xlsx", "excel", "read", "table"],
            },
            example='office.inspect_xlsx(path="report.xlsx")',
        ),
    ]

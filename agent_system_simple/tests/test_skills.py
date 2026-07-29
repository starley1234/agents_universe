"""Тесты навыков: документы, RAG, верификация, fetch, PostgreSQL.

Проверяется на НАСТОЯЩИХ файлах и настоящем HTTP-сервере. Отдельное
внимание — способности тестов ловить ошибки: для верификации ложное
«требование закрыто» опаснее пропуска, поэтому оно проверяется прямо.
"""
from __future__ import annotations

import json
import socket
import sys
import tempfile
import threading
import time
import zlib
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from agent.skills import documents as D          # noqa: E402
from agent.skills import rag as R                # noqa: E402
from agent.skills import verify as V             # noqa: E402
from agent.store import Store                    # noqa: E402
from agent.tools import fetch as F               # noqa: E402
from agent.tools.base import ToolError, Workspace  # noqa: E402

PASS, FAIL = 0, 0


def check(name: str, cond: bool, detail: str = "") -> None:
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  ok   {name}")
    else:
        FAIL += 1
        print(f"  FAIL {name}" + (f" — {detail}" if detail else ""))


def section(t: str) -> None:
    print(f"\n{t}\n" + "─" * len(t))


def free_port() -> int:
    s = socket.socket()
    s.bind(("127.0.0.1", 0))
    p = s.getsockname()[1]
    s.close()
    return p


# ══════════════════════════ документы ═══════════════════════════
def test_documents() -> None:
    section("Документы: чтение форматов")
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(td)
        t = {x.name: x for x in D.build(ws)}

        # csv
        (ws.root / "d.csv").write_text("дата,сумма\n01.02.2026,15000 руб\n",
                                       encoding="utf-8")
        out = t["doc_read"].fn(path="d.csv")
        check("csv → таблица Markdown", "| дата | сумма |" in out, out[:60])

        # docx через библиотеку, если есть
        try:
            import docx                                    # noqa: F401
            import docx as _d
            doc = _d.Document()
            doc.add_heading("Требования", 1)
            doc.add_paragraph("Изделие должно соответствовать ГОСТ 12345.")
            doc.save(ws.root / "a.docx")
            out = t["doc_read"].fn(path="a.docx")
            check("docx → заголовок сохранён", "# Требования" in out, out[:80])
            check("docx → текст извлечён", "ГОСТ 12345" in out)
        except ImportError:
            print("  skip docx (нет python-docx)")

        # xlsx
        try:
            import openpyxl
            wb = openpyxl.Workbook()
            sh = wb.active
            sh.title = "Ведомость"
            sh.append(["Поз", "Наименование"])
            sh.append(["1", "Венец"])
            wb.save(ws.root / "b.xlsx")
            out = t["doc_read"].fn(path="b.xlsx")
            check("xlsx → таблица", "| Поз | Наименование |" in out, out[:80])
            check("xlsx → данные", "Венец" in out)
        except ImportError:
            print("  skip xlsx (нет openpyxl)")

        # pdf со сжатым потоком
        content = b"BT (Trebovaniya po FAP 21) Tj ET"
        comp = zlib.compress(content)
        pdf = (b"%PDF-1.4\n4 0 obj<</Length " + str(len(comp)).encode()
               + b"/Filter/FlateDecode>>stream\n" + comp
               + b"\nendstream endobj\n%%EOF")
        (ws.root / "c.pdf").write_bytes(pdf)
        out = t["doc_read"].fn(path="c.pdf")
        check("pdf → текст из Flate-потока", "FAP 21" in out, out[:100])
        check("pdf → предупреждение об ограничениях", "разбор PDF" in out)

        # НЕГАТИВНЫЕ
        try:
            t["doc_read"].fn(path="нет.pdf")
            check("отсутствующий файл отвергнут", False)
        except ToolError:
            check("отсутствующий файл отвергнут", True)
        (ws.root / "x.doc").write_bytes(b"old binary")
        try:
            t["doc_read"].fn(path="x.doc")
            check(".doc отвергнут с подсказкой", False)
        except ToolError as exc:
            check(".doc отвергнут с подсказкой", "docx" in str(exc))


def test_documents_no_libs() -> None:
    section("Документы: работа БЕЗ сторонних библиотек")
    # Система обещает ноль обязательных зависимостей. Прячем docx и
    # openpyxl и требуем, чтобы резервный путь на stdlib работал.
    # Здесь ловится дефект inlineStr: openpyxl без sharedStrings пишет
    # строки в <is>, а не в <v>, и таблица выходила пустой.
    import builtins
    real = builtins.__import__

    def blocked(name, *a, **k):
        if name in ("docx", "openpyxl", "pypdf"):
            raise ImportError(f"{name} спрятан для теста")
        return real(name, *a, **k)

    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(td)
        try:
            import docx as _d
            import openpyxl as _x
        except ImportError:
            print("  skip (библиотек нет — резервный путь и так основной)")
            return
        doc = _d.Document()
        doc.add_heading("Требования", 1)
        doc.add_paragraph("Текст по ГОСТ 12345.")
        doc.save(ws.root / "a.docx")
        wb = _x.Workbook()
        sh = wb.active
        sh.title = "Ведомость"
        sh.append(["Поз", "Наименование"])
        sh.append(["1", "Венец"])
        wb.save(ws.root / "b.xlsx")

        builtins.__import__ = blocked
        try:
            t = {x.name: x for x in D.build(ws)}
            out = t["doc_read"].fn(path="a.docx")
            check("docx на stdlib: заголовок", "# Требования" in out, out[:70])
            check("docx на stdlib: текст", "ГОСТ 12345" in out)
            out = t["doc_read"].fn(path="b.xlsx")
            check("xlsx на stdlib: шапка", "| Поз | Наименование |" in out,
                  out[:90])
            check("xlsx на stdlib: данные (inlineStr)", "Венец" in out,
                  out[:90])
        finally:
            builtins.__import__ = real


def test_classify_extract() -> None:
    section("Документы: классификация и сущности")
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(td)
        t = {x.name: x for x in D.build(ws)}
        (ws.root / "n.md").write_text(
            "# Требования\n\nЗаявитель обязан представить документы "
            "согласно ФАП 21 и ГОСТ Р 51000-99 от 01.02.2026 № 15/АБ.\n\n"
            "## Параметры\n\nМасса 12.5 мм, отклонение 3 %.\n",
            encoding="utf-8")

        out = t["doc_classify"].fn(path="n.md")
        check("класс определён как нормативный", "нормативный" in out, out[:60])

        data = json.loads(t["doc_to_json"].fn(path="n.md"))
        ents = data["entities"]
        check("извлечена дата", "01.02.2026" in ents.get("дата", []))
        check("извлечён ГОСТ", any("51000" in g for g in ents.get("гост", [])))
        check("извлечён ФАП", "ФАП 21" in ents.get("фап", []))
        check("извлечён размер", "12.5 мм" in ents.get("размер_мм", []))
        check("разделы разобраны", len(data["sections"]) == 2,
              str(len(data["sections"])))

        # НЕГАТИВНЫЙ: пустой текст не должен рождать сущности
        (ws.root / "e.md").write_text("просто текст без ничего\n",
                                      encoding="utf-8")
        d2 = json.loads(t["doc_to_json"].fn(path="e.md"))
        check("из пустого текста сущности не выдумываются",
              not d2["entities"], str(d2["entities"]))
        check("неопознанный тип помечен как прочее",
              d2["class"] == "прочее", d2["class"])


def test_doc_graph() -> None:
    section("Документы: связи в графе")
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(td)
        st = Store(str(Path(td) / "g.db"))
        rid = st.start_run("док")
        t = {x.name: x for x in D.build(ws, st, lambda: rid)}
        (ws.root / "n.md").write_text(
            "# Раздел А\n\nСогласно ГОСТ Р 51000-99 от 01.02.2026.\n",
            encoding="utf-8")
        out = t["doc_link"].fn(path="n.md")
        check("связи созданы", "создано связей" in out, out[:60])
        nb = st.neighbours("document", "n.md")
        preds = {n["pred"] for n in nb}
        check("есть связь с разделом", "содержит_раздел" in preds, str(preds))
        check("есть связь с сущностью", "упоминает" in preds, str(preds))
        st.close()


# ═════════════════════════════ RAG ══════════════════════════════
def test_rag() -> None:
    section("RAG: обычный поиск и по онтологии")
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(td)
        d = ws.root / "docs"
        d.mkdir()
        (d / "fap21.md").write_text(
            "# ФАП 21\n\nЗаявитель обязан представить документацию.\n\n"
            "## Отклонения\n\nДопуск не более 0.1 мм.\n", encoding="utf-8")
        (d / "gost.md").write_text(
            "# ГОСТ Р 51000\n\nТребования к органам сертификации.\n",
            encoding="utf-8")
        st = Store(str(ws.root / "r.db"))
        rid = st.start_run("rag")
        t = {x.name: x for x in R.build(ws, st, lambda: rid)}

        out = t["rag_index"].fn(path="docs")
        check("документы проиндексированы", "документов: 2" in out, out[:80])

        out = t["rag_search"].fn(query="допуск отклонение")
        check("обычный поиск находит", "0.1 мм" in out, out[:120])

        out = t["rag_search"].fn(query="абракадабра квантовая")
        check("несвязанный запрос не выдумывает",
              "ничего не найдено" in out.lower(), out[:80])

        out = t["rag_search_ontology"].fn(query="ФАП")
        check("поиск по онтологии находит документ", "fap21.md" in out,
              out[:120])

        out = t["rag_answer"].fn(query="допуск")
        check("контекст содержит источник", "fap21.md" in out, out[:100])
        check("контекст требует опираться на данные",
              "ТОЛЬКО по этим данным" in out)

        out = t["rag_status"].fn()
        check("статус честен про лексический поиск",
              "эмбеддинги" in out, out[-120:])

        # НЕГАТИВНЫЙ: поиск без индекса
        st2 = Store(str(ws.root / "empty.db"))
        t2 = {x.name: x for x in R.build(ws, st2, lambda: 0)}
        try:
            t2["rag_search"].fn(query="что-нибудь")
            check("пустой индекс сообщает о себе", False)
        except ToolError as exc:
            check("пустой индекс сообщает о себе", "rag_index" in str(exc))
        st2.close()
        st.close()


# ═══════════════════════════ верификация ════════════════════════
def test_verify() -> None:
    section("Верификация: комплектность и покрытие")
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(td)
        req = ws.root / "req"
        ev = ws.root / "ev"
        req.mkdir()
        ev.mkdir()
        (req / "fap.md").write_text(
            "1.1 Заявитель должен представить доказательную документацию.\n"
            "1.2 Конструкция должна выдерживать нагрузку не менее 500 кг.\n"
            "1.3 Не допускается применение материалов без сертификата.\n"
            "1.4 Изделие должно пройти климатические испытания.\n",
            encoding="utf-8")
        (ev / "protokol.md").write_text(
            "# Протокол № 12/25 от 01.02.2026\nУтверждаю. Исполнитель: Иванов.\n"
            "Испытания нагрузкой: конструкция выдержала нагрузку 500 кг.\n"
            "Масса изделия 12.5 кг.\n", encoding="utf-8")
        (ev / "sert.md").write_text(
            "# Сертификат № 7/АБ от 15.01.2026\nПодписано. Разработчик: Альфа.\n"
            "Масса изделия 13.0 кг.\n", encoding="utf-8")
        t = {x.name: x for x in V.build(ws)}

        # ГЛАВНОЕ: все требования должны находиться, включая короткие
        out = t["verify_requirements"].fn(req_path="req", evidence_path="ev")
        check("найдены ВСЕ 4 требования", "Найдено требований: 4" in out,
              out[:200])

        # ГЛАВНОЕ: климатические испытания НЕ должны считаться закрытыми
        # протоколом нагрузки — ложное «закрыто» опаснее пропуска
        clim_line = next((l for l in out.split("\n")
                          if "климатические" in l), "")
        check("климатические испытания не в «ЗАКРЫТО»",
              not clim_line.strip().startswith("+"), clim_line[:80])
        check("покрытие честное, не 100%", "(25%)" in out or "(50%)" in out,
              [l for l in out.split("\n") if "закрыто:" in l.lower()])

        # комплектность
        out = t["check_completeness"].fn(path="ev",
                                         required="протокол, сертификат, "
                                                  "климатические испытания")
        check("недостающий документ найден", "ОТСУТСТВУЕТ" in out
              and "климатические" in out, out[:150])
        check("комплект помечен неполным", "НЕ полон" in out)

        # реквизиты
        (ev / "bez.md").write_text("Просто текст без реквизитов.\n",
                                   encoding="utf-8")
        out = t["check_requisites"].fn(path="ev")
        check("документ без реквизитов найден", "bez.md: НЕТ" in out,
              out[:200])

        # сверка чисел
        out = t["cross_check"].fn(path="ev")
        check("расхождение массы найдено", "12.5 кг" in out and "13.0 кг" in out,
              out[:200])

        # оговорка обязана быть в каждом отчёте
        for name, res in (("верификация",
                           t["verify_requirements"].fn(req_path="req",
                                                       evidence_path="ev")),
                          ("комплектность",
                           t["check_completeness"].fn(path="ev",
                                                      required="протокол")),
                          ("сверка", t["cross_check"].fn(path="ev"))):
            check(f"{name}: есть оговорка об эксперте",
                  "эксперт" in res.lower(), res[-100:])

        # сводный отчёт
        out = t["audit_report"].fn(req_path="req", evidence_path="ev",
                                   required="протокол, сертификат")
        check("сводный отчёт создан", "audit.md" in out, out[:80])
        rep = (ws.root / "audit.md").read_text(encoding="utf-8")
        check("в отчёте все четыре раздела",
              all(x in rep for x in ("1. Комплектность", "2. Реквизиты",
                                     "3. Покрытие", "4. Согласованность")))
        check("в отчёте есть дисклеймер",
              "не является заключением" in rep, rep[-200:])


# ═══════════════════════════════ fetch ══════════════════════════
HTML = (b"<html><head><title>Test Page</title></head><body>"
        b"<script>var x=1;</script><nav>menu</nav>"
        b"<h1>Zagolovok</h1><p>Osnovnoy tekst stranicy.</p>"
        b'<a href="/doc1">Pervaya</a><a href="https://ext.example/x">Vtoraya</a>'
        b"</body></html>")


class _H(BaseHTTPRequestHandler):
    def log_message(self, *a):
        pass

    def do_GET(self):
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(HTML)))
        self.end_headers()
        self.wfile.write(HTML)


def test_fetch() -> None:
    section("Fetch: загрузка страниц и защита сети")
    port = free_port()
    srv = ThreadingHTTPServer(("127.0.0.1", port), _H)
    threading.Thread(target=srv.serve_forever, daemon=True).start()
    time.sleep(0.3)
    try:
        with tempfile.TemporaryDirectory() as td:
            ws = Workspace(td)
            url = f"http://127.0.0.1:{port}/"

            # ГЛАВНОЕ: внутренняя сеть закрыта по умолчанию
            t = {x.name: x for x in F.build(ws)}
            try:
                t["fetch_url"].fn(url=url)
                check("localhost заблокирован", False, "агент вышел в LAN!")
            except ToolError as exc:
                check("localhost заблокирован", "внутренний" in str(exc))

            t = {x.name: x for x in F.build(ws, allow_private=True)}
            out = t["fetch_url"].fn(url=url)
            check("страница загружена", "Osnovnoy tekst" in out, out[:100])
            check("заголовок извлечён", "Test Page" in out)
            check("скрипты вычищены", "var x" not in out)
            check("указан источник", url in out)

            out = t["fetch_links"].fn(url=url)
            check("ссылки собраны", "ext.example" in out, out[:100])
            check("относительная ссылка развёрнута", "/doc1" in out)

            out = t["fetch_url"].fn(url=url, save_to="page.html")
            check("сохранение работает", (ws.root / "page.html").exists(), out)

            # НЕГАТИВНЫЕ
            for bad, why in ((f"ftp://127.0.0.1:{port}/", "схема"),
                             ("не-урл-вовсе", "хост")):
                try:
                    t["fetch_url"].fn(url=bad)
                    check(f"отклонён {why}", False, bad)
                except ToolError:
                    check(f"отклонён {why}", True)
    finally:
        srv.shutdown()


# ════════════════════════════ PostgreSQL ════════════════════════
def test_pg_graceful() -> None:
    section("PostgreSQL: недоступность не ломает агента")
    from agent.skills import pgonto as PG
    pg = PG.PgOnto(dsn="postgresql://nobody@127.0.0.1:1/nodb")
    ok = pg.connect()
    check("подключение не удалось (ожидаемо)", not ok)
    check("причина сохранена", bool(pg.error), pg.error[:60])

    t = {x.name: x for x in PG.build(pg)}
    out = t["pg_status"].fn()
    check("статус объясняет проблему", "НЕДОСТУПЕН" in out, out[:80])
    check("статус подсказывает про SQLite", "SQLite" in out)

    try:
        t["pg_add_entity"].fn(kind="part", name="венец")
        check("операция без БД отвергнута", False, "прошла!")
    except ToolError as exc:
        check("операция без БД отвергнута", "PostgreSQL" in str(exc))

    # эмбеддинги не настроены — честный отказ, а не пустой результат
    try:
        t["pg_semantic_search"].fn(query="что-нибудь")
        check("поиск без эмбеддингов отвергнут", False)
    except ToolError as exc:
        check("поиск без эмбеддингов отвергнут",
              "PostgreSQL" in str(exc) or "эмбеддинг" in str(exc))


def test_build_profiles() -> None:
    section("Сборка агентов по профилям")
    from agent.build import build_agent, known_skills
    from agent.config import Config

    for s in ("documents", "rag", "verify", "pg", "fetch"):
        check(f"навык {s} зарегистрирован", s in known_skills())

    for prof, must in (("docs", "doc_to_json"), ("rag", "rag_search_ontology"),
                       ("verify", "verify_requirements"),
                       ("onto", "pg_semantic_search")):
        with tempfile.TemporaryDirectory() as td:
            cfg = Config.load(None, profile=prof, provider="ollama",
                              model="m", workspace=td)
            cfg.db = str(Path(td) / "a.db")
            agent = build_agent(cfg)
            check(f"профиль {prof}: есть {must}", must in agent.tools.names(),
                  str(len(agent.tools)))


def main() -> int:
    print("=" * 60)
    print("ТЕСТЫ НАВЫКОВ: документы, RAG, верификация, fetch, PostgreSQL")
    print("=" * 60)
    test_documents()
    test_documents_no_libs()
    test_classify_extract()
    test_doc_graph()
    test_rag()
    test_verify()
    test_fetch()
    test_pg_graceful()
    test_build_profiles()
    print("\n" + "=" * 60)
    print(f"пройдено: {PASS} · провалено: {FAIL}")
    print("=" * 60)
    return 1 if FAIL else 0


if __name__ == "__main__":
    raise SystemExit(main())

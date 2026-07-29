"""Чтение документов: PDF, Word, Excel, CSV, текст → Markdown / JSON.

Принцип системы — ноль обязательных зависимостей, поэтому все форматы
разбираются стандартной библиотекой:

  docx/xlsx  это ZIP с XML внутри  → zipfile + ElementTree
  pdf        потоки, сжатые zlib   → zlib + разбор текстовых операторов
  csv/txt    напрямую

Если в системе ОКАЖУТСЯ python-docx / openpyxl / pypdf — используем их,
качество разбора выше. Но работа не ломается, когда их нет.

Честная граница: собственный PDF-парсер берёт текст из несжатых и
Flate-сжатых потоков. Он НЕ читает сканы (там картинки, нужен OCR) и
может путать порядок колонок в сложной вёрстке. Об этом сообщается
прямо в выводе, а не замалчивается.
"""
from __future__ import annotations

import csv
import io
import json
import re
import zipfile
import zlib
from pathlib import Path
from typing import Any

from ..tools.base import Tool, ToolError, Workspace

MAX_CHARS = 200_000          # предел текста на документ
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
S_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"


# ────────────────────────────── PDF ──────────────────────────────
def _pdf_streams(raw: bytes) -> list[bytes]:
    """Содержимое всех потоков документа, по возможности распакованное."""
    out: list[bytes] = []
    for m in re.finditer(rb"stream\r?\n(.*?)endstream", raw, re.S):
        data = m.group(1)
        try:
            out.append(zlib.decompress(data))
        except zlib.error:
            out.append(data)          # поток без сжатия
    return out


def _pdf_text(raw: bytes) -> tuple[str, str]:
    """Текст PDF и примечание об ограничениях."""
    # если есть готовая библиотека — она точнее
    try:
        import pypdf                                    # type: ignore
        rd = pypdf.PdfReader(io.BytesIO(raw))
        txt = "\n\n".join((p.extract_text() or "") for p in rd.pages)
        return txt.strip(), ""
    except Exception:
        pass

    chunks: list[str] = []
    for st in _pdf_streams(raw):
        # (текст) Tj   и   [(a) -2 (b)] TJ
        for m in re.finditer(rb"\((?:\\.|[^\\()])*\)", st):
            s = m.group(0)[1:-1]
            s = (s.replace(rb"\(", b"(").replace(rb"\)", b")")
                  .replace(rb"\\", b"\\"))
            if s.strip():
                chunks.append(s.decode("utf-8", "replace"))
    text = " ".join(chunks)
    text = re.sub(r"[ \t]{2,}", " ", text).strip()
    note = ""
    if not text:
        note = ("текст не извлечён: вероятно скан (нужен OCR) или "
                "нестандартная кодировка шрифта")
    else:
        note = ("разбор PDF стандартной библиотекой: порядок блоков в "
                "сложной вёрстке может отличаться от визуального")
    return text, note


# ────────────────────────────── DOCX ─────────────────────────────
def _docx_text(path: Path) -> str:
    """Абзацы и таблицы Word в Markdown."""
    try:
        import docx                                     # type: ignore
        d = docx.Document(str(path))
        parts = []
        for p in d.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            lvl = 0
            if p.style is not None and p.style.name.startswith("Heading"):
                m = re.search(r"\d+", p.style.name)
                lvl = int(m.group()) if m else 1
            parts.append(("#" * min(lvl, 6) + " " + t) if lvl else t)
        for tb in d.tables:
            parts.append(_rows_to_md([[c.text.strip() for c in r.cells]
                                      for r in tb.rows]))
        return "\n\n".join(parts)
    except ImportError:
        pass

    # запасной путь: docx это ZIP с word/document.xml
    import xml.etree.ElementTree as ET
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    parts = []
    for p in root.iter(f"{W_NS}p"):
        txt = "".join(t.text or "" for t in p.iter(f"{W_NS}t")).strip()
        if not txt:
            continue
        style = p.find(f"{W_NS}pPr/{W_NS}pStyle")
        lvl = 0
        if style is not None:
            v = style.get(f"{W_NS}val", "")
            m = re.search(r"(\d+)", v)
            if "eading" in v and m:
                lvl = int(m.group(1))
        parts.append(("#" * min(lvl, 6) + " " + txt) if lvl else txt)
    return "\n\n".join(parts)


# ────────────────────────────── XLSX ─────────────────────────────
def _xlsx_tables(path: Path, max_rows: int = 500) -> dict[str, list[list[str]]]:
    try:
        import openpyxl                                 # type: ignore
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        out = {}
        for ws in wb.worksheets:
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    break
                rows.append(["" if c is None else str(c) for c in row])
            out[ws.title] = rows
        wb.close()
        return out
    except ImportError:
        pass

    # запасной путь: xlsx это ZIP; строки лежат в общей таблице строк
    import xml.etree.ElementTree as ET
    with zipfile.ZipFile(path) as z:
        shared: list[str] = []
        if "xl/sharedStrings.xml" in z.namelist():
            sr = ET.fromstring(z.read("xl/sharedStrings.xml"))
            shared = ["".join(t.text or "" for t in si.iter(f"{S_NS}t"))
                      for si in sr.iter(f"{S_NS}si")]
        names = {}
        if "xl/workbook.xml" in z.namelist():
            wbr = ET.fromstring(z.read("xl/workbook.xml"))
            for i, sh in enumerate(wbr.iter(f"{S_NS}sheet"), 1):
                names[i] = sh.get("name", f"Лист{i}")
        out = {}
        for n in sorted(x for x in z.namelist()
                        if re.fullmatch(r"xl/worksheets/sheet\d+\.xml", x)):
            idx = int(re.search(r"(\d+)", n.split("/")[-1]).group(1))
            root = ET.fromstring(z.read(n))
            rows = []
            for r in list(root.iter(f"{S_NS}row"))[:max_rows]:
                cells = []
                for c in r.iter(f"{S_NS}c"):
                    # Excel хранит строки ТРЕМЯ способами, и учитывать надо
                    # все: иначе часть ячеек молча выходит пустой.
                    t = c.get("t")
                    if t == "inlineStr":            # <is><t>текст</t></is>
                        node = c.find(f"{S_NS}is")
                        val = ("".join(x.text or "" for x in node.iter(f"{S_NS}t"))
                               if node is not None else "")
                    else:
                        v = c.find(f"{S_NS}v")
                        val = "" if v is None else (v.text or "")
                        if t == "s" and val.isdigit():   # общая таблица строк
                            i = int(val)
                            val = shared[i] if i < len(shared) else ""
                    cells.append(val)
                rows.append(cells)
            out[names.get(idx, f"Лист{idx}")] = rows
        return out


def _rows_to_md(rows: list[list[str]]) -> str:
    rows = [r for r in rows if any(str(c).strip() for c in r)]
    if not rows:
        return ""
    w = max(len(r) for r in rows)
    rows = [list(r) + [""] * (w - len(r)) for r in rows]
    head = "| " + " | ".join(rows[0]) + " |"
    sep = "|" + "|".join(["---"] * w) + "|"
    body = ["| " + " | ".join(str(c) for c in r) + " |" for r in rows[1:]]
    return "\n".join([head, sep, *body])


# ───────────────────────── единая точка входа ────────────────────
def read_any(path: Path) -> tuple[str, dict[str, Any]]:
    """Документ → (markdown, метаданные). Формат по расширению."""
    ext = path.suffix.lower()
    meta: dict[str, Any] = {"file": path.name, "format": ext.lstrip("."),
                            "bytes": path.stat().st_size}
    if ext == ".pdf":
        text, note = _pdf_text(path.read_bytes())
        if note:
            meta["note"] = note
        return text, meta
    if ext == ".docx":
        return _docx_text(path), meta
    if ext in (".xlsx", ".xlsm"):
        tables = _xlsx_tables(path)
        meta["sheets"] = list(tables)
        parts = [f"## {name}\n\n{_rows_to_md(rows)}"
                 for name, rows in tables.items() if rows]
        return "\n\n".join(parts), meta
    if ext == ".csv":
        raw = path.read_text(encoding="utf-8", errors="replace")
        dialect = csv.Sniffer().sniff(raw[:2000]) if raw.strip() else csv.excel
        rows = list(csv.reader(io.StringIO(raw), dialect))
        meta["rows"] = len(rows)
        return _rows_to_md(rows[:500]), meta
    if ext in (".md", ".txt", ".json", ".xml", ".html", ".log", ".yml", ".yaml"):
        return path.read_text(encoding="utf-8", errors="replace"), meta
    if ext == ".doc":
        raise ToolError(
            ".doc — устаревший двоичный формат Word. Пересохраните в .docx "
            "(в Word: Файл → Сохранить как) или конвертируйте: "
            "libreoffice --headless --convert-to docx файл.doc")
    raise ToolError(f"формат {ext!r} не поддержан. Доступны: pdf, docx, "
                    f"xlsx, csv, md, txt, json, xml, html")


# ───────────────────────── классификация ─────────────────────────
#: Ключевые слова по типам документов. Список открытый: неизвестный тип
#: помечается как "прочее", а не подгоняется под ближайший.
CLASSES: dict[str, list[str]] = {
    "нормативный": ["требование", "должен", "обязан", "не допускается",
                    "устанавливает", "регламентирует", "стандарт", "ГОСТ",
                    "СП ", "приказ", "федеральные авиационные правила", "ФАП"],
    "технический": ["чертёж", "спецификация", "материал", "допуск", "расчёт",
                    "конструкция", "прочность", "нагрузка", "мм", "кг"],
    "договор": ["стороны", "исполнитель", "заказчик", "обязуется",
                "настоящий договор", "оплата", "приложение №"],
    "отчёт": ["результаты", "выводы", "проведено", "испытани", "измерени",
              "заключение", "протокол"],
    "заявка": ["заявитель", "прошу", "заявка", "выдать", "сертификат"],
    "переписка": ["уважаемый", "письмо", "исх. №", "вх. №", "направляем"],
}


def classify(text: str) -> tuple[str, dict[str, int]]:
    low = text.lower()
    scores = {k: sum(low.count(w.lower()) for w in words)
              for k, words in CLASSES.items()}
    best = max(scores, key=lambda k: scores[k])
    return (best if scores[best] >= 2 else "прочее"), scores


#: Извлечение сущностей. Только то, что распознаётся ОДНОЗНАЧНО —
#: догадки в структурированные данные попадать не должны.
PATTERNS: dict[str, str] = {
    "дата": r"\b\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4}\b",
    "номер_документа": r"№\s?[\w\-/.]+",
    "гост": r"\bГОСТ\s?[\dР\-. ]+\d",
    "фап": r"\bФАП[\s\-]?\d+\b",
    "сумма": r"\b\d[\d\s]{2,}(?:[.,]\d{2})?\s?(?:руб|₽|RUB)\b",
    "размер_мм": r"\b\d+(?:[.,]\d+)?\s?мм\b",
    "процент": r"\b\d+(?:[.,]\d+)?\s?%",
    "email": r"\b[\w.\-]+@[\w.\-]+\.\w{2,}\b",
    "url": r"https?://[^\s)>\]]+",
}


def extract(text: str) -> dict[str, list[str]]:
    out: dict[str, list[str]] = {}
    for name, rx in PATTERNS.items():
        found = re.findall(rx, text, flags=re.I)
        uniq = list(dict.fromkeys(f.strip() for f in found))
        if uniq:
            out[name] = uniq[:40]
    return out


def sections(md: str) -> list[dict[str, Any]]:
    """Разбиение по заголовкам Markdown — основа для связей и RAG."""
    out: list[dict[str, Any]] = []
    cur = {"title": "", "level": 0, "text": []}
    for line in md.split("\n"):
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            if cur["text"] or cur["title"]:
                out.append({"title": cur["title"], "level": cur["level"],
                            "text": "\n".join(cur["text"]).strip()})
            cur = {"title": m.group(2).strip(), "level": len(m.group(1)),
                   "text": []}
        else:
            cur["text"].append(line)
    if cur["text"] or cur["title"]:
        out.append({"title": cur["title"], "level": cur["level"],
                    "text": "\n".join(cur["text"]).strip()})
    return [s for s in out if s["text"] or s["title"]]


# ──────────────────────────── инструменты ────────────────────────
def build(ws: Workspace, store=None, run_id_getter=None) -> list[Tool]:

    def doc_read(path: str, max_chars: int = 0) -> str:
        p = ws.resolve(path)
        if not p.exists():
            raise ToolError(f"файл {path!r} не найден")
        text, meta = read_any(p)
        lim = max_chars if max_chars > 0 else MAX_CHARS
        head = f"# {meta['file']}\n\n"
        if meta.get("note"):
            head += f"> {meta['note']}\n\n"
        if meta.get("sheets"):
            head += f"> листов: {', '.join(meta['sheets'])}\n\n"
        if len(text) > lim:
            text = text[:lim] + f"\n\n… обрезано, всего {len(text)} символов"
        return head + (text or "(документ пуст или текст не извлечён)")

    def doc_to_json(path: str, out_path: str = "") -> str:
        p = ws.resolve(path)
        if not p.exists():
            raise ToolError(f"файл {path!r} не найден")
        text, meta = read_any(p)
        kind, scores = classify(text)
        data = {
            "source": meta["file"],
            "format": meta["format"],
            "bytes": meta["bytes"],
            "class": kind,
            "class_scores": {k: v for k, v in scores.items() if v},
            "entities": extract(text),
            "sections": [{"title": s["title"], "level": s["level"],
                          "chars": len(s["text"]),
                          "text": s["text"][:4000]} for s in sections(text)],
            "chars": len(text),
        }
        if meta.get("note"):
            data["note"] = meta["note"]
        js = json.dumps(data, ensure_ascii=False, indent=2)
        if out_path:
            o = ws.resolve(out_path)
            o.parent.mkdir(parents=True, exist_ok=True)
            o.write_text(js, encoding="utf-8")
            return (f"JSON сохранён: {ws.relative(o)}\n"
                    f"класс: {kind}, разделов: {len(data['sections'])}, "
                    f"типов сущностей: {len(data['entities'])}")
        return js[:MAX_CHARS]

    def doc_to_markdown(path: str, out_path: str = "") -> str:
        p = ws.resolve(path)
        text, meta = read_any(p)
        md = f"# {meta['file']}\n\n"
        if meta.get("note"):
            md += f"> {meta['note']}\n\n"
        md += text
        if out_path:
            o = ws.resolve(out_path)
            o.parent.mkdir(parents=True, exist_ok=True)
            o.write_text(md, encoding="utf-8")
            return f"Markdown сохранён: {ws.relative(o)} ({len(md)} символов)"
        return md[:MAX_CHARS]

    def doc_classify(path: str) -> str:
        p = ws.resolve(path)
        text, _ = read_any(p)
        kind, scores = classify(text)
        top = sorted(scores.items(), key=lambda x: -x[1])[:4]
        lines = [f"Класс: {kind}", "", "Совпадений по признакам:"]
        lines += [f"  {k}: {v}" for k, v in top if v]
        ents = extract(text)
        if ents:
            lines += ["", "Найденные сущности:"]
            lines += [f"  {k}: {', '.join(v[:6])}"
                      + (f" … ещё {len(v)-6}" if len(v) > 6 else "")
                      for k, v in ents.items()]
        if kind == "прочее":
            lines += ["", "Признаков известных классов мало — "
                          "тип определить не удалось."]
        return "\n".join(lines)

    def doc_link(path: str, kind: str = "document") -> str:
        """Документ → объекты и связи в графе знаний."""
        if store is None:
            raise ToolError("нужен навык memory: граф знаний не подключён")
        p = ws.resolve(path)
        text, meta = read_any(p)
        cls, _ = classify(text)
        ents = extract(text)
        rid = run_id_getter() if run_id_getter else 0
        doc = meta["file"]
        store.upsert_entity(kind, doc, {"class": cls, "format": meta["format"],
                                        "chars": len(text)}, run_id=rid)
        n = 0
        for sec in sections(text):
            if not sec["title"]:
                continue
            if store.link((kind, doc), "содержит_раздел",
                          ("раздел", sec["title"][:80]), run_id=rid):
                n += 1
        for etype, vals in ents.items():
            for v in vals[:10]:
                if store.link((kind, doc), "упоминает", (etype, v[:80]),
                              run_id=rid):
                    n += 1
        e, r = store.graph_stats()
        return (f"{doc}: класс «{cls}», создано связей {n}\n"
                f"В графе: {e} объектов, {r} связей")

    return [
        Tool("doc_read",
             "Прочитать документ любого формата (PDF, Word .docx, Excel "
             ".xlsx, CSV, txt, md) и получить его текст в Markdown.",
             {"type": "object",
              "properties": {"path": {"type": "string"},
                             "max_chars": {"type": "integer"}},
              "required": ["path"]},
             doc_read),
        Tool("doc_to_json",
             "Разобрать документ в структурированный JSON: класс, разделы, "
             "найденные сущности (даты, номера, ГОСТ, суммы). Без out_path "
             "возвращает JSON текстом.",
             {"type": "object",
              "properties": {"path": {"type": "string"},
                             "out_path": {"type": "string"}},
              "required": ["path"]},
             doc_to_json),
        Tool("doc_to_markdown",
             "Конвертировать документ в Markdown с сохранением заголовков "
             "и таблиц.",
             {"type": "object",
              "properties": {"path": {"type": "string"},
                             "out_path": {"type": "string"}},
              "required": ["path"]},
             doc_to_markdown),
        Tool("doc_classify",
             "Определить тип документа (нормативный, технический, договор, "
             "отчёт, заявка, переписка) и извлечь сущности.",
             {"type": "object",
              "properties": {"path": {"type": "string"}},
              "required": ["path"]},
             doc_classify),
        Tool("doc_link",
             "Разобрать документ и занести его в граф знаний: сам документ, "
             "его разделы и упомянутые сущности со связями.",
             {"type": "object",
              "properties": {"path": {"type": "string"},
                             "kind": {"type": "string"}},
              "required": ["path"]},
             doc_link),
    ]

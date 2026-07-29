"""Навык «docparse»: чтение Word/Excel/текстовых файлов, определение типа
раздела/листа, преобразование в markdown/JSON и выделение связанных
объектов онтологии — тот же конвейер идей, что у навыка pdf, но без
рендера в картинку: docx/xlsx уже дают точный текстовый/табличный слой,
поэтому извлечение делается ДЕТЕРМИНИРОВАННО (без LLM, без риска, что
модель ошибётся в цифре), а модель подключается только там, где нужно
подлинное понимание естественного языка — при выделении сущностей.

Три шага:

  ШАГ 1 (doc_info, doc_classify) — как и в pdf: локально, без LLM.
  Документ Word делится на РАЗДЕЛЫ (последовательность абзацев/таблиц
  между заголовками), документ Excel — на ЛИСТЫ; для каждого куска
  считаются те же признаки (доля текста, наличие таблицы, ключевые
  слова BOM/счёта), что и для страниц PDF, и подставляются в ОБЩЕЕ
  дерево решений doc_types.classify_signals — поэтому лист Excel с
  колонками «Поз/Обозначение/Кол-во» получит тот же тип bom_table, что
  и аналогичная таблица в PDF.

  ШАГ 2 (doc_extract) — конвертация в markdown/JSON. Для bom_table и
  invoice_financial таблица сериализуется в JSON ПРЯМО из ячеек (никакой
  LLM, 100% точность чисел). Для остальных типов текст переносится в
  markdown с сохранением заголовков/абзацев/таблиц.

  ШАГ 3 (extract_entities) — берёт готовый markdown/JSON (неважно, чей:
  этого навыка или navyка pdf) и текстовой LLM выделяет сущности и связи
  предметной области, которые тут же пишутся в общую онтологию
  (Store.entity/relation) — то самое «создавать из них связанные
  объекты». Проверка на выдумывание частично встроена в промпт
  (doc_types.ENTITY_SYSTEM), но окончательная ответственность — на
  модели; агент должен относиться к результату как к черновику онтологии,
  а не как к истине в последней инстанции.

python-docx и openpyxl — единственные внешние зависимости, импортируются
ЛЕНИВО, как pymupdf в навыке pdf.
"""
from __future__ import annotations

import csv
import json
import re
from pathlib import Path
from typing import Any

from . import doc_types as dt
from ..llm.base import BaseLLM, LLMError
from ..store import Store
from ..tools.base import Tool, ToolError, Workspace

TEXT_EXTS = {".txt", ".md", ".markdown"}
CSV_EXTS = {".csv", ".tsv"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx", ".xlsm"}
SUPPORTED_EXTS = TEXT_EXTS | CSV_EXTS | DOCX_EXTS | XLSX_EXTS


def _require_docx():
    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise ToolError(
            "Чтение .docx требует python-docx. Установите: pip install python-docx"
        ) from exc
    return docx


def _require_openpyxl():
    try:
        import openpyxl  # type: ignore
    except ImportError as exc:
        raise ToolError(
            "Чтение .xlsx требует openpyxl. Установите: pip install openpyxl"
        ) from exc
    return openpyxl


# ------------------------------------------------------ docx: разбор на разделы
def _iter_docx_blocks(doc, docx_mod):
    """Абзацы и таблицы В ПОРЯДКЕ, в котором они идут в документе.

    python-docx по умолчанию отдаёт paragraphs и tables отдельными
    списками без порядка — так заголовок и следующая за ним таблица
    расходятся. Обходим сырое XML-дерево тела документа.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _table_to_rows(table) -> list[list[str]]:
    return [[c.text.strip() for c in row.cells] for row in table.rows]


def _heading_level(style_name: str) -> int:
    """0 = Title, 1..9 = Heading N, -1 = не заголовок."""
    if style_name == "Title":
        return 0
    m = re.match(r"Heading (\d+)", style_name or "")
    return int(m.group(1)) if m else -1


class DocxSection:
    """Раздел docx: заголовок + всё содержимое до следующего заголовка
    того же или более высокого уровня."""

    def __init__(self, title: str, level: int) -> None:
        self.title = title
        self.level = level
        self.paragraphs: list[str] = []
        self.tables: list[list[list[str]]] = []
        self.max_font = 0.0     # приближённо: заголовок считаем крупным
        self.median_font = 12.0

    def signals(self) -> dict[str, Any]:
        text = "\n".join(self.paragraphs)
        table_text = " ".join(
            " ".join(cell for row in t for cell in row) for t in self.tables)
        full = text + " " + table_text
        return {
            "chars": len(text.strip()),
            "words": len(text.split()),
            "tables": len(self.tables),
            "vector_paths": 0,          # в docx векторной графики не считаем
            "image_area_ratio": 0.0,
            "landscape": False,
            "max_font": 16.0 if self.level == 0 else 12.0,
            "median_font": 12.0,
            "bom_hits": dt.keyword_hits(full, dt.BOM_WORDS),
            "invoice_hits": dt.keyword_hits(full, dt.INVOICE_WORDS),
            "drawing_hits": dt.keyword_hits(full, dt.DRAWING_WORDS),
        }

    def to_markdown(self) -> str:
        parts = []
        if self.title:
            parts.append("#" * max(1, self.level or 1) + " " + self.title)
        for p in self.paragraphs:
            if p.strip():
                parts.append(p)
        for t in self.tables:
            if not t:
                continue
            header, *rows = t
            parts.append("| " + " | ".join(header) + " |")
            parts.append("|" + "---|" * len(header))
            for row in rows:
                parts.append("| " + " | ".join(row) + " |")
        return "\n\n".join(parts)


def _split_docx_sections(path: Path, docx_mod) -> list[DocxSection]:
    doc = docx_mod.Document(str(path))
    sections: list[DocxSection] = [DocxSection("", 1)]  # преамбула до первого заголовка
    for block in _iter_docx_blocks(doc, docx_mod):
        cls_name = type(block).__name__
        if cls_name == "Paragraph":
            lvl = _heading_level(block.style.name if block.style else "")
            if lvl >= 0 and block.text.strip():
                sections.append(DocxSection(block.text.strip(), max(1, lvl)))
            elif block.text.strip():
                sections[-1].paragraphs.append(block.text.strip())
        else:  # Table
            sections[-1].tables.append(_table_to_rows(block))
    # преамбула без содержимого — шум, убираем
    if sections and not sections[0].title and not sections[0].paragraphs \
            and not sections[0].tables:
        sections = sections[1:]
    return sections or [DocxSection("", 1)]


# ------------------------------------------------------- xlsx: разбор листов
def _sheet_signals(rows: list[list[Any]]) -> dict[str, Any]:
    flat = " ".join(str(c) for row in rows for c in row if c is not None)
    non_empty_rows = [r for r in rows if any(c is not None and str(c).strip()
                                             for c in r)]
    has_table = len(non_empty_rows) >= 2 and rows and len(rows[0]) >= 2
    return {
        "chars": len(flat.strip()),
        "words": len(flat.split()),
        "tables": 1 if has_table else 0,
        "vector_paths": 0,
        "image_area_ratio": 0.0,
        "landscape": False,
        "max_font": 12.0,
        "median_font": 12.0,
        "bom_hits": dt.keyword_hits(flat, dt.BOM_WORDS),
        "invoice_hits": dt.keyword_hits(flat, dt.INVOICE_WORDS),
        "drawing_hits": dt.keyword_hits(flat, dt.DRAWING_WORDS),
    }


def _sheet_rows(ws) -> list[list[Any]]:
    return [list(row) for row in ws.iter_rows(values_only=True)]


def _rows_to_markdown_table(rows: list[list[Any]]) -> str:
    if not rows:
        return "(лист пуст)"
    header = [str(c) if c is not None else "" for c in rows[0]]
    out = ["| " + " | ".join(header) + " |", "|" + "---|" * len(header)]
    for row in rows[1:]:
        cells = [str(c) if c is not None else "" for c in row]
        cells += [""] * (len(header) - len(cells))
        out.append("| " + " | ".join(cells[:len(header)]) + " |")
    return "\n".join(out)


def _rows_to_bom_json(rows: list[list[Any]], sheet: str) -> dict[str, Any]:
    if not rows:
        return {"type": "bom_table", "sheet": sheet, "columns": [], "items": []}
    columns = [str(c) if c is not None else f"col{i}"
              for i, c in enumerate(rows[0])]
    items = []
    for row in rows[1:]:
        if not any(c is not None and str(c).strip() for c in row):
            continue
        item = {columns[i]: (row[i] if i < len(row) else None)
                for i in range(len(columns))}
        items.append(item)
    return {"type": "bom_table", "sheet": sheet, "columns": columns, "items": items}


# --------------------------------------------------------------- csv/txt
def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _csv_rows(path: Path, delimiter: str) -> list[list[str]]:
    with path.open(newline="", encoding="utf-8", errors="replace") as fh:
        return [row for row in csv.reader(fh, delimiter=delimiter)]


def _strip_code_fence(text: str) -> str:
    t = text.strip()
    m = re.match(r"^```(?:json|markdown|md)?\s*\n(.*)\n```$", t, re.S)
    return m.group(1).strip() if m else t


# ================================================================== build
def build(ws: Workspace, text_llm: BaseLLM, store: Store | None,
          run_id_getter) -> list[Tool]:

    def _resolve(path: str) -> Path:
        p = ws.resolve(path)
        if not p.exists():
            raise ToolError(f"Файл {path!r} не найден")
        return p

    def _ext(p: Path) -> str:
        return p.suffix.lower()

    def _check_supported(p: Path) -> None:
        if _ext(p) == ".pdf":
            raise ToolError(
                "Для PDF используйте навык pdf (pdf_info/pdf_classify/"
                "pdf_extract) — этот навык обрабатывает Word/Excel/текст."
            )
        if _ext(p) not in SUPPORTED_EXTS:
            raise ToolError(
                f"Формат {_ext(p) or '(без расширения)'} не поддержан. "
                f"Поддержаны: {', '.join(sorted(SUPPORTED_EXTS))}"
            )

    # --------------------------------------------------------- doc_info
    def doc_info(path: str) -> str:
        p = _resolve(path)
        _check_supported(p)
        ext = _ext(p)
        if ext in DOCX_EXTS:
            docx_mod = _require_docx()
            secs = _split_docx_sections(p, docx_mod)
            lines = [f"{ws.relative(p)}: Word, {len(secs)} раздел(ов)"]
            for i, s in enumerate(secs, 1):
                lines.append(f"  раздел {i}: {s.title or '(без заголовка)'} — "
                            f"{len(s.paragraphs)} абз., {len(s.tables)} табл.")
            return "\n".join(lines)
        if ext in XLSX_EXTS:
            openpyxl = _require_openpyxl()
            wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
            lines = [f"{ws.relative(p)}: Excel, {len(wb.sheetnames)} лист(ов)"]
            for name in wb.sheetnames:
                sh = wb[name]
                lines.append(f"  лист {name!r}: {sh.max_row} строк, "
                            f"{sh.max_column} столбцов")
            wb.close()
            return "\n".join(lines)
        if ext in CSV_EXTS:
            rows = _csv_rows(p, "\t" if ext == ".tsv" else ",")
            return f"{ws.relative(p)}: {ext[1:].upper()}, {len(rows)} строк"
        text = _read_text(p)
        return (f"{ws.relative(p)}: текст, {len(text)} симв., "
                f"{len(text.splitlines())} строк")

    # ------------------------------------------------------ doc_classify
    def doc_classify(path: str) -> str:
        """ШАГ 1: тип каждого раздела Word / листа Excel / текста целиком."""
        p = _resolve(path)
        _check_supported(p)
        ext = _ext(p)
        out = [f"{ws.relative(p)}: классификация"]
        if ext in DOCX_EXTS:
            docx_mod = _require_docx()
            secs = _split_docx_sections(p, docx_mod)
            for i, s in enumerate(secs, 1):
                doc_type, conf, reasons = dt.classify_signals(s.signals())
                out.append(f"  раздел {i} ({s.title or 'без заголовка'}): "
                          f"{doc_type} (увер. {conf:.2f}) — {'; '.join(reasons)}")
        elif ext in XLSX_EXTS:
            openpyxl = _require_openpyxl()
            wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
            for name in wb.sheetnames:
                rows = _sheet_rows(wb[name])
                doc_type, conf, reasons = dt.classify_signals(_sheet_signals(rows))
                out.append(f"  лист {name!r}: {doc_type} (увер. {conf:.2f}) — "
                          f"{'; '.join(reasons)}")
            wb.close()
        elif ext in CSV_EXTS:
            rows = _csv_rows(p, "\t" if ext == ".tsv" else ",")
            doc_type, conf, reasons = dt.classify_signals(_sheet_signals(rows))
            out.append(f"  таблица: {doc_type} (увер. {conf:.2f}) — "
                      f"{'; '.join(reasons)}")
        else:
            text = _read_text(p)
            sig = {"chars": len(text.strip()), "words": len(text.split()),
                  "tables": 0, "vector_paths": 0, "image_area_ratio": 0.0,
                  "landscape": False, "max_font": 12.0, "median_font": 12.0,
                  "bom_hits": dt.keyword_hits(text, dt.BOM_WORDS),
                  "invoice_hits": dt.keyword_hits(text, dt.INVOICE_WORDS),
                  "drawing_hits": dt.keyword_hits(text, dt.DRAWING_WORDS)}
            doc_type, conf, reasons = dt.classify_signals(sig)
            out.append(f"  текст: {doc_type} (увер. {conf:.2f}) — "
                      f"{'; '.join(reasons)}")
        out.append("Доступные типы: " + ", ".join(
            f"{k} ({v})" for k, v in dt.DOC_TYPES.items()))
        return "\n".join(out)

    # -------------------------------------------------------- doc_extract
    def doc_extract(path: str, out_format: str = "auto",
                    save_to: str = "") -> str:
        """ШАГ 2: детерминированное преобразование в markdown/JSON.

        Для табличных типов (bom_table/invoice_financial) отдаём точный
        JSON из ячеек — ни одна цифра не проходит через модель и не
        может быть искажена распознаванием.
        """
        p = _resolve(path)
        _check_supported(p)
        ext = _ext(p)
        pieces: list[dict[str, Any]] = []   # [{unit, doc_type, format, text}]

        if ext in DOCX_EXTS:
            docx_mod = _require_docx()
            secs = _split_docx_sections(p, docx_mod)
            for i, s in enumerate(secs, 1):
                doc_type, _, _ = dt.classify_signals(s.signals())
                fmt = out_format
                if fmt == "auto":
                    fmt = "json" if dt.looks_like_json(doc_type) and s.tables \
                        else "markdown"
                if fmt == "json" and s.tables:
                    text = json.dumps(
                        _rows_to_bom_json(s.tables[0], s.title or f"раздел {i}"),
                        ensure_ascii=False, indent=2)
                else:
                    fmt = "markdown"
                    text = s.to_markdown()
                pieces.append({"unit": s.title or f"раздел {i}",
                               "doc_type": doc_type, "format": fmt, "text": text})
        elif ext in XLSX_EXTS:
            openpyxl = _require_openpyxl()
            wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
            for name in wb.sheetnames:
                rows = _sheet_rows(wb[name])
                doc_type, _, _ = dt.classify_signals(_sheet_signals(rows))
                fmt = out_format
                if fmt == "auto":
                    fmt = "json" if dt.looks_like_json(doc_type) else "markdown"
                if fmt == "json":
                    text = json.dumps(_rows_to_bom_json(rows, name),
                                      ensure_ascii=False, indent=2)
                else:
                    fmt = "markdown"
                    text = f"## {name}\n\n" + _rows_to_markdown_table(rows)
                pieces.append({"unit": name, "doc_type": doc_type,
                              "format": fmt, "text": text})
            wb.close()
        elif ext in CSV_EXTS:
            rows = _csv_rows(p, "\t" if ext == ".tsv" else ",")
            doc_type, _, _ = dt.classify_signals(_sheet_signals(rows))
            fmt = out_format
            if fmt == "auto":
                fmt = "json" if dt.looks_like_json(doc_type) else "markdown"
            text = (json.dumps(_rows_to_bom_json(rows, p.name), ensure_ascii=False,
                               indent=2) if fmt == "json"
                    else _rows_to_markdown_table(rows))
            pieces.append({"unit": p.name, "doc_type": doc_type,
                          "format": fmt, "text": text})
        else:
            text = _read_text(p)
            sig = {"chars": len(text.strip()), "words": len(text.split()),
                  "tables": 0, "vector_paths": 0, "image_area_ratio": 0.0,
                  "landscape": False, "max_font": 12.0, "median_font": 12.0,
                  "bom_hits": 0, "invoice_hits": 0, "drawing_hits": 0}
            doc_type, _, _ = dt.classify_signals(sig)
            pieces.append({"unit": p.name, "doc_type": doc_type,
                          "format": "markdown", "text": text})

        lines = [f"{ws.relative(p)}: обработано {len(pieces)} фрагмент(ов)"]
        combined: list[str] = []
        for piece in pieces:
            header = f"[{piece['unit']}, тип {piece['doc_type']}, " \
                    f"формат {piece['format']}]"
            lines.append(f"  {header}: {len(piece['text'])} симв.")
            combined.append(header + "\n" + piece["text"])
        full_text = "\n\n---\n\n".join(combined)
        if save_to.strip():
            out_p = ws.resolve(save_to)
            out_p.parent.mkdir(parents=True, exist_ok=True)
            out_p.write_text(full_text, encoding="utf-8")
            lines.append(f"Сохранено: {ws.relative(out_p)}")
        else:
            lines.append("\n" + full_text)
        return "\n".join(lines)

    # -------------------------------------------------------- extract_entities
    def extract_entities(text: str = "", path: str = "", doc_type: str = "",
                         source: str = "") -> str:
        """ШАГ 3: выделить сущности и связи из ГОТОВОГО текста (markdown/
        JSON — неважно, от pdf_extract или doc_extract) и записать их в
        общую онтологию. Текст можно передать напрямую или файлом (path).
        """
        if path.strip():
            p = ws.resolve(path)
            if not p.exists():
                raise ToolError(f"Файл {path!r} не найден")
            body = p.read_text(encoding="utf-8", errors="replace")
            src = source or ws.relative(p)
        elif text.strip():
            body = text
            src = source or "(инлайн-текст)"
        else:
            raise ToolError("Нужно указать либо text, либо path")

        dtype = doc_type.strip() or "mixed_report"
        if dtype not in dt.DOC_TYPES:
            raise ToolError(
                f"Неизвестный doc_type {dtype!r}. Доступны: "
                f"{', '.join(dt.DOC_TYPES)}"
            )

        prompt = dt.ENTITY_PROMPT_TEMPLATE.format(
            doc_type=dtype, doc_type_desc=dt.DOC_TYPES[dtype],
            source=src, text=body[:12000])
        try:
            reply = text_llm.chat([
                {"role": "system", "content": dt.ENTITY_SYSTEM},
                {"role": "user", "content": prompt},
            ], tools=None)
        except LLMError as exc:
            raise ToolError(f"Модель извлечения сущностей недоступна: {exc}") from exc

        raw = _strip_code_fence(reply.text or "")
        try:
            data = json.loads(raw)
        except json.JSONDecodeError as exc:
            raise ToolError(
                f"Модель вернула не-JSON при выделении сущностей: {exc}. "
                f"Ответ: {raw[:300]}"
            ) from exc

        entities = data.get("entities") or []
        relations = data.get("relations") or []
        if not isinstance(entities, list) or not isinstance(relations, list):
            raise ToolError("Ответ модели должен содержать списки entities/relations")

        if store is None:
            return (f"Разобрано (без сохранения — навык memory не подключён): "
                    f"{len(entities)} сущностей, {len(relations)} связей")

        rid = run_id_getter() if run_id_getter else None
        n_ent = 0
        for e in entities:
            kind, name = str(e.get("kind", "")).strip(), str(e.get("name", "")).strip()
            if not kind or not name:
                continue
            store.upsert_entity(kind, name, e.get("props") or {}, run_id=rid)
            n_ent += 1
        n_rel = 0
        for r in relations:
            subj, obj = r.get("subject"), r.get("object")
            pred = str(r.get("predicate", "")).strip()
            if not (isinstance(subj, list) and len(subj) == 2 and
                    isinstance(obj, list) and len(obj) == 2 and pred):
                continue
            store.link((str(subj[0]), str(subj[1])), pred,
                      (str(obj[0]), str(obj[1])), run_id=rid)
            n_rel += 1

        e_total, r_total = store.graph_stats()
        return (f"Из {src!r} ({dtype}) сохранено: {n_ent} объектов, "
               f"{n_rel} связей. Всего в онтологии: {e_total} объектов, "
               f"{r_total} связей.")

    tools = [
        Tool("doc_info",
             "Быстрая сводка по Word/Excel/CSV/текстовому файлу: разделы/"
             "листы и их размер. Смотреть перед классификацией/извлечением "
             "на большом файле. Для PDF используйте pdf_info.",
             {"type": "object",
              "properties": {"path": {"type": "string"}},
              "required": ["path"]},
             doc_info),
        Tool("doc_classify",
             "ШАГ 1: определить тип каждого раздела Word / листа Excel / "
             "CSV-таблицы / текстового файла — теми же правилами, что и "
             "pdf_classify (bom_table, invoice_financial, prose_text и т.д.), "
             "чтобы решить, как извлекать содержимое на шаге 2.",
             {"type": "object",
              "properties": {"path": {"type": "string"}},
              "required": ["path"]},
             doc_classify),
        Tool("doc_extract",
             "ШАГ 2: детерминированно (без LLM, без риска ошибиться в "
             "цифрах) конвертировать Word/Excel/CSV/текст в markdown или "
             "JSON с учётом типа каждого раздела/листа. Табличные типы "
             "(bom_table/invoice_financial) отдаются строгим JSON прямо "
             "из ячеек.",
             {"type": "object",
              "properties": {
                  "path": {"type": "string"},
                  "out_format": {"type": "string",
                                 "description": "auto|markdown|json"},
                  "save_to": {"type": "string",
                              "description": "Куда сохранить результат "
                                            "(если пусто — вернуть текстом)"}},
              "required": ["path"]},
             doc_extract),
        Tool("extract_entities",
             "ШАГ 3: выделить объекты предметной области и связи между "
             "ними из УЖЕ ГОТОВОГО текста (markdown/JSON, полученного "
             "pdf_extract_page/pdf_extract или doc_extract) и записать их "
             "в общую онтологию (те же таблицы, что у note_entity/link). "
             "Так классифицированный документ превращается в связанные "
             "объекты графа знаний. Передайте text ИЛИ path.",
             {"type": "object",
              "properties": {
                  "text": {"type": "string", "description": "Текст напрямую"},
                  "path": {"type": "string", "description": "Или файл с текстом"},
                  "doc_type": {"type": "string",
                               "description": "Тип документа: " +
                                             ", ".join(dt.DOC_TYPES)},
                  "source": {"type": "string",
                             "description": "Имя источника для пометки объектов"}},
              "required": []},
             extract_entities),
    ]
    return tools

"""Тесты навыка docparse: чтение Word/Excel/CSV/текста, классификация,
детерминированное извлечение и выделение сущностей в онтологию.

Требуют python-docx и openpyxl. Если их нет — тесты пропускаются с
понятным сообщением, а не падают с ImportError посреди прогона.

Философия та же, что у test_pdf.py: тест обязан уметь падать. Рядом с
позитивными проверками — негативные: неподдержанный формат, PDF через
неверный навык, отсутствующий файл, невалидный ответ модели.
"""
from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from agent.llm.base import BaseLLM, LLMReply, Usage       # noqa: E402
from agent.store import Store                             # noqa: E402
from agent.tools.base import ToolError, Workspace         # noqa: E402

PASS, FAIL = 0, 0

try:
    import docx  # type: ignore
    import openpyxl  # type: ignore
    HAVE_DEPS = True
except ImportError:
    HAVE_DEPS = False


def check(name: str, cond: bool, detail: str = "") -> None:
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  ok   {name}")
    else:
        FAIL += 1
        print(f"  FAIL {name}" + (f" — {detail}" if detail else ""))


def section(title: str) -> None:
    print(f"\n{title}\n" + "─" * len(title))


# ------------------------------------------------------------- фикстуры
def _make_bom_docx(path: Path) -> None:
    d = docx.Document()
    d.add_heading("Титульный документ", 0)
    d.add_paragraph("Некоторое вступление для контекста, чтобы был связный "
                    "текст достаточной длины для классификации как прозы.")
    d.add_heading("Спецификация", 1)
    t = d.add_table(rows=3, cols=3)
    t.cell(0, 0).text, t.cell(0, 1).text, t.cell(0, 2).text = \
        "Поз", "Обозначение", "Наименование"
    t.cell(1, 0).text, t.cell(1, 1).text, t.cell(1, 2).text = "1", "AB-01", "Корпус"
    t.cell(2, 0).text, t.cell(2, 1).text, t.cell(2, 2).text = "2", "AB-02", "Крышка"
    d.save(path)


def _make_bom_xlsx(path: Path) -> None:
    wb = openpyxl.Workbook()
    sh = wb.active
    sh.title = "BOM"
    sh.append(["Поз", "Обозначение", "Наименование", "Кол"])
    sh.append([1, "AB-01", "Корпус", 1])
    sh.append([2, "AB-02", "Крышка", 1])
    sh2 = wb.create_sheet("Notes")
    sh2["A1"] = "Заметка без структуры, просто текст без таблицы."
    wb.save(path)


def _import_docparse():
    from agent.skills import docparse
    return docparse


# --------------------------------------------------------- фальшивая LLM
class FakeTextLLM(BaseLLM):
    """Отдаёт заранее заданный JSON-ответ для extract_entities."""

    def __init__(self, text: str = "") -> None:
        super().__init__("fake-text")
        self.text = text
        self.calls: list[list[dict]] = []

    def chat(self, messages, tools=None):
        self.calls.append(list(messages))
        return LLMReply(text=self.text, usage=Usage(20, 10))


# ================================================================ tests
def test_docx_classify_and_extract() -> None:
    section("Word: разбиение на разделы, классификация, извлечение")
    docparse = _import_docparse()
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(Path(td) / "ws")
        _make_bom_docx(ws.root / "doc.docx")
        tools = {t.name: t for t in docparse.build(ws, FakeTextLLM(), None,
                                                    lambda: 0)}

        info = tools["doc_info"].fn(path="doc.docx")
        check("doc_info видит 2 раздела", "2 раздел" in info, info)

        cls = tools["doc_classify"].fn(path="doc.docx")
        check("титульный раздел не bom_table", "раздел 1" in cls and
              "bom_table" not in cls.split("раздел 1")[1].split("раздел 2")[0])
        check("раздел со спецификацией -> bom_table",
              "раздел 2" in cls and "bom_table" in cls.split("раздел 2")[1])

        out = tools["doc_extract"].fn(path="doc.docx")
        check("markdown для титульного раздела", "формат markdown" in out)
        check("JSON для спецификации", "формат json" in out)
        # JSON-фрагмент — после заголовка [Спецификация, ...] до конца вывода
        marker = "[Спецификация, тип bom_table, формат json]\n"
        json_text = out.split(marker, 1)[1]
        parsed = json.loads(json_text)
        check("JSON распарсился", parsed["type"] == "bom_table")
        check("значения ячеек переданы точно",
              parsed["items"][0]["Обозначение"] == "AB-01")



def test_xlsx_classify_and_extract() -> None:
    section("Excel: классификация листов, точный JSON из ячеек")
    docparse = _import_docparse()
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(Path(td) / "ws")
        _make_bom_xlsx(ws.root / "book.xlsx")
        tools = {t.name: t for t in docparse.build(ws, FakeTextLLM(), None,
                                                    lambda: 0)}

        cls = tools["doc_classify"].fn(path="book.xlsx")
        check("лист BOM -> bom_table", "'BOM': bom_table" in cls, cls)
        check("лист Notes не bom_table",
              "'Notes': bom_table" not in cls)

        out = tools["doc_extract"].fn(path="book.xlsx", out_format="json")
        # оба листа принудительно json при out_format=json
        check("оба листа обработаны", "обработано 2 фрагмент" in out)


def test_csv_and_text() -> None:
    section("CSV и обычный текст")
    docparse = _import_docparse()
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(Path(td) / "ws")
        (ws.root / "data.csv").write_text(
            "Поз,Обозначение,Наименование\n1,AB-01,Корпус\n2,AB-02,Крышка\n",
            encoding="utf-8")
        (ws.root / "note.txt").write_text(
            "Обычный текстовый файл без всякой структуры, просто заметка "
            "для последующего чтения человеком когда-нибудь потом.",
            encoding="utf-8")
        tools = {t.name: t for t in docparse.build(ws, FakeTextLLM(), None,
                                                    lambda: 0)}

        cls_csv = tools["doc_classify"].fn(path="data.csv")
        check("CSV с BOM-заголовками -> bom_table", "bom_table" in cls_csv, cls_csv)

        out_csv = tools["doc_extract"].fn(path="data.csv")
        check("CSV извлечён как JSON", "формат json" in out_csv)

        cls_txt = tools["doc_classify"].fn(path="note.txt")
        check("текстовый файл классифицирован", "текст:" in cls_txt, cls_txt)


def test_extract_entities_writes_ontology() -> None:
    section("extract_entities: сущности и связи пишутся в общую онтологию")
    docparse = _import_docparse()
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(Path(td) / "ws")
        store = Store(str(Path(td) / "test.db"))
        fake = FakeTextLLM(text=json.dumps({
            "entities": [
                {"kind": "part", "name": "AB-01", "props": {"material": "steel"}},
                {"kind": "assembly", "name": "Редуктор", "props": {}},
            ],
            "relations": [
                {"subject": ["part", "AB-01"], "predicate": "входит_в",
                 "object": ["assembly", "Редуктор"]},
            ],
        }))
        tools = {t.name: t for t in docparse.build(ws, fake, store, lambda: 0)}

        out = tools["extract_entities"].fn(
            text="AB-01 — корпус, входит в Редуктор.", doc_type="bom_table",
            source="doc.docx")
        check("отчёт содержит число сущностей", "2 объектов" in out, out)
        check("отчёт содержит число связей", "1 связей" in out, out)

        e, r = store.graph_stats()
        check("сущности реально в базе", e == 2, str(e))
        check("связь реально в базе", r == 1, str(r))
        check("промпт содержит имя источника",
              "doc.docx" in fake.calls[-1][-1]["content"])


def test_negative_cases() -> None:
    section("Негативные проверки")
    docparse = _import_docparse()
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(Path(td) / "ws")
        _make_bom_docx(ws.root / "doc.docx")
        tools = {t.name: t for t in docparse.build(ws, FakeTextLLM(), None,
                                                    lambda: 0)}

        try:
            tools["doc_info"].fn(path="nope.docx")
            check("отказ на отсутствующий файл", False)
        except ToolError:
            check("отказ на отсутствующий файл", True)

        (ws.root / "image.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        try:
            tools["doc_info"].fn(path="image.png")
            check("отказ на неподдержанный формат", False)
        except ToolError as exc:
            check("отказ на неподдержанный формат", True)
            check("сообщение перечисляет поддержанные форматы",
                  ".docx" in str(exc), str(exc))

        (ws.root / "fake.pdf").write_bytes(b"%PDF-1.4 not really")
        try:
            tools["doc_info"].fn(path="fake.pdf")
            check("PDF направляется к другому навыку", False)
        except ToolError as exc:
            check("PDF направляется к другому навыку", "навык pdf" in str(exc), str(exc))

        try:
            tools["extract_entities"].fn(text="", path="")
            check("отказ без text и path", False)
        except ToolError:
            check("отказ без text и path", True)

        try:
            tools["extract_entities"].fn(text="что-то", doc_type="выдуманный")
            check("отказ на неизвестный doc_type", False)
        except ToolError:
            check("отказ на неизвестный doc_type", True)

        bad_llm = FakeTextLLM(text="это не json, а просто текст")
        tools2 = {t.name: t for t in docparse.build(ws, bad_llm, None, lambda: 0)}
        try:
            tools2["extract_entities"].fn(text="что-то про AB-01")
            check("отказ при невалидном JSON от модели", False)
        except ToolError:
            check("отказ при невалидном JSON от модели", True)


def test_extract_entities_without_store() -> None:
    section("extract_entities без подключённого store (только разбор)")
    docparse = _import_docparse()
    with tempfile.TemporaryDirectory() as td:
        ws = Workspace(Path(td) / "ws")
        fake = FakeTextLLM(text=json.dumps({"entities": [], "relations": []}))
        tools = {t.name: t for t in docparse.build(ws, fake, None, lambda: 0)}
        out = tools["extract_entities"].fn(text="что-то")
        check("сообщение о том, что не сохранено", "без сохранения" in out, out)


def test_build_agent_with_docparse_skill() -> None:
    section("Сборка агента с навыком docparse")
    from agent.build import build_agent
    from agent.config import Config
    with tempfile.TemporaryDirectory() as td:
        cfg = Config(provider="ollama", model="qwen2.5-coder", workspace=td,
                    skills=["files", "docparse"])
        agent = build_agent(cfg)
        names = agent.tools.names()
        for t in ("doc_info", "doc_classify", "doc_extract", "extract_entities"):
            check(f"инструмент {t} зарегистрирован", t in names)


def main() -> int:
    if not HAVE_DEPS:
        print("python-docx/openpyxl не установлены — тесты docparse "
              "пропущены (pip install python-docx openpyxl)")
        return 0
    test_docx_classify_and_extract()
    test_xlsx_classify_and_extract()
    test_csv_and_text()
    test_extract_entities_writes_ontology()
    test_negative_cases()
    test_extract_entities_without_store()
    test_build_agent_with_docparse_skill()

    print(f"\n{'─' * 40}\nитого: {PASS} ok, {FAIL} fail")
    return 1 if FAIL else 0


if __name__ == "__main__":
    raise SystemExit(main())
